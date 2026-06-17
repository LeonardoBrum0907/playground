#!/usr/bin/env python3
"""Gera avaliação de Geografia (3º ano) em formato Word (.docx).

Conteúdos alinhados ao livro didático: festas populares (frevo), etnia e
cultura, elementos culturais e cultura através das gerações.
"""

import time
import urllib.request
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

OUTPUT_PATH = Path("/workspace/Avaliacao_Geografia_3ano.docx")
ASSETS_DIR = Path("/workspace/avaliacao_geografia_assets")
GABARITO_PATH = Path("/workspace/Gabarito_Avaliacao_Geografia_3ano.docx")

GREEN = RGBColor(27, 94, 32)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)

IMAGE_SOURCES = {
    "frevo": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/b/b0/"
        "Frevo_no_Carnaval_do_Recife.jpg/960px-Frevo_no_Carnaval_do_Recife.jpg"
    ),
    "etnia_mundo": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/4/4c/"
        "West_Greenland_Inuit_modern_outfit_with_avittat.jpg/960px-"
        "West_Greenland_Inuit_modern_outfit_with_avittat.jpg"
    ),
    "cultura_elementos": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/7/74/"
        "Music_is_international_language_210824-F-F3261-1001.jpg/960px-"
        "Music_is_international_language_210824-F-F3261-1001.jpg"
    ),
    "esconde_esconde": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/2/22/"
        "Crian%C3%A7as_CEU.jpg/960px-Crian%C3%A7as_CEU.jpg"
    ),
    "cultura_costumes": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/9/92/"
        "Three_generations_of_a_family_%281%29.jpg/960px-"
        "Three_generations_of_a_family_%281%29.jpg"
    ),
    "cultura_diversidade": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e8/"
        "VanAchterbergAlgeria.jpg/960px-VanAchterbergAlgeria.jpg"
    ),
    "cultura_urbana": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3b/"
        "Grafite_no_BankBoston_da_Avenida_Paulista%2C_S%C3%A3o_Paulo.jpg/"
        "960px-Grafite_no_BankBoston_da_Avenida_Paulista%2C_S%C3%A3o_Paulo.jpg"
    ),
}

OBJETIVAS = [
    {
        "number": 1,
        "theme": "Festas populares e frevo",
        "question": (
            "Observe a imagem. A dança representada é o frevo, uma festa popular "
            "de rua. Sobre essa manifestação cultural, assinale a alternativa correta:"
        ),
        "image": "frevo",
        "caption": "Passistas de frevo no Carnaval do Recife, Pernambuco.",
        "options": [
            "O frevo é uma dança que se originou em Pernambuco e faz parte das festas populares brasileiras.",
            "O frevo é uma dança que existe apenas em países europeus.",
            "O frevo não é considerado manifestação cultural porque acontece na rua.",
            "O frevo é igual em todas as cidades do mundo.",
        ],
        "answer": "A",
    },
    {
        "number": 2,
        "theme": "Etnia e cultura",
        "question": (
            "Segundo o livro, etnia é uma comunidade formada por pessoas que "
            "compartilham elementos comuns. Observe a imagem e marque a alternativa "
            "correta sobre cultura:"
        ),
        "image": "etnia_mundo",
        "caption": "Traje tradicional inuit — exemplo de identidade cultural.",
        "options": [
            "Todas as culturas são iguais e não apresentam diferenças.",
            "Algumas culturas são melhores que outras.",
            "Existem culturas diferentes, e nenhuma é melhor ou pior — são apenas diferentes.",
            "A cultura existe apenas nas grandes cidades.",
        ],
        "answer": "C",
    },
    {
        "number": 3,
        "theme": "O que é cultura",
        "question": (
            "O livro apresenta crenças, expressões artísticas, costumes e leis como "
            "parte da cultura. Qual alternativa representa uma expressão artística?"
        ),
        "image": "cultura_elementos",
        "caption": "Crianças cantando e tocando instrumento musical.",
        "options": [
            "Respeitar as leis de trânsito ao atravessar a rua.",
            "Almoçar em família ao redor da mesa.",
            "Cantar e tocar instrumentos musicais.",
            "Acreditar em ideias e tradições religiosas.",
        ],
        "answer": "C",
    },
    {
        "number": 4,
        "theme": "Cultura através das gerações",
        "question": (
            "Observe a imagem. Qual brincadeira está sendo representada e como ela "
            "se relaciona com a cultura?"
        ),
        "image": "esconde_esconde",
        "caption": "Crianças brincando ao ar livre — transmissão cultural entre gerações.",
        "options": [
            "Amarelinha; é uma brincadeira que não muda nunca.",
            "Esconde-esconde; é uma brincadeira aprendida e transmitida entre pessoas de diferentes idades.",
            "Videogame; só existe na cultura dos adultos.",
            "Futebol; não faz parte da cultura infantil.",
        ],
        "answer": "B",
    },
]

DISSERTATIVAS = [
    {
        "number": 5,
        "theme": "Festas populares",
        "question": (
            "O que é uma festa popular? Explique usando o frevo como exemplo. "
            "Informe em qual estado brasileiro essa dança se originou e descreva "
            "dois elementos visíveis na imagem (como roupa, sombrinha ou movimento)."
        ),
        "image": "frevo",
        "caption": "Frevo — manifestação cultural de Pernambuco.",
        "lines": 6,
        "answer_hint": (
            "Festa popular é uma celebração realizada pelo povo, muitas vezes na "
            "rua. O frevo é originário de Pernambuco (Recife/Olinda). Elementos: "
            "fantasia colorida, sombrinha (guarda-sol), passos acrobáticos, música."
        ),
    },
    {
        "number": 6,
        "theme": "Etnia e cultura",
        "question": (
            "Escreva com suas palavras o que é etnia. Depois, observe a imagem e "
            "cite dois elementos culturais que podem ser identificados na roupa e "
            "no modo de vida das pessoas fotografadas."
        ),
        "image": "etnia_mundo",
        "caption": "Povos de diferentes regiões têm trajes e costumes próprios.",
        "lines": 6,
        "answer_hint": (
            "Etnia é um grupo de pessoas que compartilha cultura, língua, história "
            "e características comuns. Elementos: traje tradicional, padrões "
            "têxteis, adaptação ao clima, formas de trabalho e convivência."
        ),
    },
    {
        "number": 7,
        "theme": "O que é cultura",
        "question": (
            "De acordo com o livro, cite os quatro elementos que formam a cultura "
            "de um grupo (crenças, expressões artísticas, costumes e leis) e dê "
            "um exemplo de cada um."
        ),
        "image": "cultura_costumes",
        "caption": "Família reunida — exemplo de costume cultural.",
        "lines": 7,
        "answer_hint": (
            "Crenças: religião ou ideias; Expressões artísticas: música e dança; "
            "Costumes: almoço em família; Leis: respeitar faixa de pedestres."
        ),
    },
    {
        "number": 8,
        "theme": "Cultura através das gerações",
        "question": (
            "Observe a imagem das crianças brincando. Qual brincadeira está "
            "representada? Explique como aprendemos novas brincadeiras e por que "
            "elas fazem parte da cultura transmitida entre gerações."
        ),
        "image": "esconde_esconde",
        "caption": "Brincadeiras de rua fazem parte da cultura infantil.",
        "lines": 6,
        "answer_hint": (
            "Esconde-esconde. Aprendemos com amigos, colegas e pessoas mais velhas "
            "que nos ensinam as regras. Brincadeiras são manifestações culturais "
            "passadas de geração em geração, embora também mudem com o tempo."
        ),
    },
    {
        "number": 9,
        "theme": "Cultura através das gerações",
        "question": (
            "O livro diz que, em cidades, vilas, fazendas ou mata, todos somos "
            "brasileiros, mas cada comunidade tem sua cultura. Explique por que "
            "comunidades diferentes têm costumes e tradições distintos. Dê dois "
            "exemplos de manifestações culturais (como formas de plantar, preparar "
            "comida ou brincar)."
        ),
        "image": None,
        "caption": None,
        "lines": 7,
        "answer_hint": (
            "Cada lugar tem história, clima, pessoas e convivência próprios. "
            "Exemplos: modos de plantar e preparar comida regionais; brincadeiras "
            "como pião, amarelinha, esconde-esconde; festas como junina e frevo."
        ),
    },
    {
        "number": 10,
        "theme": "Cultura brasileira",
        "question": (
            "Observe a imagem de pessoas com trajes tradicionais de outra região "
            "do mundo. Compare com a cultura brasileira: o que é semelhante e o que "
            "é diferente? Por que o livro afirma que nos sentimos pertencentes a "
            "uma comunidade?"
        ),
        "image": "cultura_diversidade",
        "caption": "Traje tradicional tuareg — diversidade cultural no mundo.",
        "lines": 7,
        "answer_hint": (
            "Semelhanças: todos têm trajes, festas, comidas e formas de viver. "
            "Diferenças: roupas, clima, trabalho e tradições variam. Sentimos "
            "pertencimento porque compartilhamos língua, costumes e identidade "
            "com o grupo em que convivemos."
        ),
    },
]


def download_images():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    headers = {"User-Agent": "GeoAssessmentBot/1.0 (educational; classroom use)"}
    paths = {}
    for name, url in IMAGE_SOURCES.items():
        path = ASSETS_DIR / f"{name}.jpg"
        downloaded = False
        for attempt in range(3):
            try:
                req = urllib.request.Request(url, headers=headers)
                with urllib.request.urlopen(req, timeout=30) as response:
                    data = response.read()
                if len(data) < 15000:
                    raise ValueError("Arquivo muito pequeno")
                path.write_bytes(data)
                paths[name] = path
                downloaded = True
                break
            except Exception:
                time.sleep(1.5 * (attempt + 1))
        if not downloaded and path.exists() and path.stat().st_size >= 15000:
            paths[name] = path
        elif not downloaded:
            raise RuntimeError(f"Não foi possível baixar a imagem: {name}")
        time.sleep(0.3)
    return paths


def set_run_font(run, size=11, bold=False, color=DARK, name="Arial"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_horizontal_line(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B8C0")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_paragraph(doc, text="", size=11, bold=False, color=DARK, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_answer_lines(doc, count):
    for _ in range(count):
        p = add_paragraph(doc, space_after=0)
        add_horizontal_line(p)
        p.paragraph_format.space_after = Pt(14)


def add_image_block(doc, image_key, caption):
    image_path = ASSETS_DIR / f"{image_key}.jpg"
    if not image_path.exists():
        add_paragraph(doc, f"[Imagem: {caption}]", size=10, color=MUTED)
        return

    with Image.open(image_path) as img:
        img = img.convert("RGB")
        buffer = BytesIO()
        img.save(buffer, format="JPEG", quality=90)
        buffer.seek(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buffer, width=Cm(11))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run_cap = cap.add_run(caption)
    set_run_font(run_cap, size=9, color=MUTED, bold=False)


def add_objective_question(doc, item):
    add_paragraph(
        doc,
        f"Questão {item['number']} — {item['theme']}",
        size=11,
        bold=True,
        color=GREEN,
        space_after=4,
    )
    add_paragraph(doc, item["question"], size=11, space_after=8)
    add_image_block(doc, item["image"], item["caption"])

    mark_table = doc.add_table(rows=1, cols=5)
    mark_table.autofit = True
    headers = ["Questão", "A", "B", "C", "D"]
    for idx, label in enumerate(headers):
        cell = mark_table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, size=10, bold=True, color=MUTED if idx else GREEN)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mark_row = mark_table.add_row().cells
    mark_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mark_row[0].paragraphs[0].add_run(f"({item['number']})")
    set_run_font(run, size=10, bold=True)
    for idx in range(1, 5):
        mark_row[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = mark_row[idx].paragraphs[0].add_run("(   )")
        set_run_font(run, size=12)

    doc.add_paragraph()

    labels = ["A", "B", "C", "D"]
    for label, option in zip(labels, item["options"]):
        p = add_paragraph(doc, space_after=4)
        run_mark = p.add_run("(   )  ")
        set_run_font(run_mark, size=11, bold=True)
        run_text = p.add_run(f"{label}) {option}")
        set_run_font(run_text, size=11)

    doc.add_paragraph()


def add_essay_question(doc, item):
    add_paragraph(
        doc,
        f"Questão {item['number']} — {item['theme']}",
        size=11,
        bold=True,
        color=GREEN,
        space_after=4,
    )
    add_paragraph(doc, item["question"], size=11, space_after=8)
    if item.get("image"):
        add_image_block(doc, item["image"], item["caption"])
    add_paragraph(doc, "Resposta:", size=11, bold=True, color=MUTED, space_after=6)
    add_answer_lines(doc, item["lines"])
    doc.add_paragraph()


def build_student_document(output_path):
    download_images()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Avaliação de Geografia")
    set_run_font(run, size=18, bold=True, color=GREEN)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("3º Ano do Ensino Fundamental")
    set_run_font(run, size=13, bold=True, color=DARK)

    theme = doc.add_paragraph()
    theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = theme.add_run(
        "Festas populares • Etnia e cultura • Elementos culturais • Cultura através das gerações"
    )
    set_run_font(run, size=10, color=MUTED)
    theme.paragraph_format.space_after = Pt(12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = info.add_run(
        "Aluno(a): ________________________________________________     "
        "Data: ____/____/________     Turma: __________"
    )
    set_run_font(run, size=11)
    info.paragraph_format.space_after = Pt(14)

    add_paragraph(doc, "Instruções", size=12, bold=True, color=GREEN, space_after=6)
    instructions = [
        "• Leia cada questão com atenção antes de responder.",
        "• Observe as imagens — elas remetem aos conteúdos estudados no livro.",
        "• Parte I: marque com X a alternativa correta nos espaços (   ).",
        "• Parte II: escreva a resposta completa nas linhas indicadas.",
        "• Dificuldade média — valor total: 10 pontos (1 ponto por questão).",
        "• Tempo sugerido: 50 minutos.",
    ]
    for line in instructions:
        add_paragraph(doc, line, size=10, space_after=3)
    doc.add_paragraph()

    add_paragraph(doc, "PARTE I — QUESTÕES OBJETIVAS", size=13, bold=True, color=GREEN, space_after=10)
    for item in OBJETIVAS:
        add_objective_question(doc, item)

    doc.add_page_break()
    add_paragraph(doc, "PARTE II — QUESTÕES DISSERTATIVAS", size=13, bold=True, color=GREEN, space_after=10)
    for item in DISSERTATIVAS:
        add_essay_question(doc, item)

    doc.save(output_path)


def build_answer_key(output_path):
    doc = Document()
    add_paragraph(doc, "Gabarito — Avaliação de Geografia (3º ano)", size=16, bold=True, color=GREEN, space_after=8)
    add_paragraph(doc, "Uso exclusivo do(a) professor(a).", size=10, color=MUTED, space_after=12)

    add_paragraph(doc, "Parte I — Objetivas", size=12, bold=True, color=GREEN, space_after=6)
    for item in OBJETIVAS:
        add_paragraph(
            doc,
            f"Questão {item['number']}: alternativa {item['answer']}",
            size=11,
            space_after=4,
        )

    doc.add_paragraph()
    add_paragraph(doc, "Parte II — Dissertativas (respostas esperadas)", size=12, bold=True, color=GREEN, space_after=6)
    for item in DISSERTATIVAS:
        add_paragraph(doc, f"Questão {item['number']}:", size=11, bold=True, space_after=2)
        add_paragraph(doc, item["answer_hint"], size=10, space_after=8)

    doc.save(output_path)


if __name__ == "__main__":
    build_student_document(OUTPUT_PATH)
    build_answer_key(GABARITO_PATH)
    print(f"Avaliação gerada: {OUTPUT_PATH}")
    print(f"Gabarito gerado: {GABARITO_PATH}")
